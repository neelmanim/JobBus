"""
JobBus Backend — Resume Analyzer.

Extracts text from PDF/DOCX and uses Gemini AI to parse
the resume into a structured profile.
"""

import io
import re
import json
import httpx
from typing import Optional

from pypdf import PdfReader

from models.schemas import ResumeProfile
from database import get_supabase_admin


class UnsupportedFormat(Exception):
    pass


class EmptyResume(Exception):
    pass


class ResumeExtractor:
    """Extracts raw text from PDF and DOCX files."""

    SUPPORTED_FORMATS = {".pdf", ".docx"}

    @staticmethod
    def extract_text(file_path: str = None, file_bytes: bytes = None, filename: str = "") -> str:
        """Extract text from a resume file.

        Args:
            file_path: Path to the file on disk.
            file_bytes: Raw bytes of the uploaded file.
            filename: Original filename (for format detection).

        Returns:
            Extracted text content.
        """
        ext = ""
        if file_path:
            ext = "." + file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
        elif filename:
            ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext not in ResumeExtractor.SUPPORTED_FORMATS:
            raise UnsupportedFormat(
                f"Unsupported format '{ext}'. Please upload a PDF or DOCX file."
            )

        if ext == ".pdf":
            return ResumeExtractor._extract_pdf(file_path, file_bytes)
        elif ext == ".docx":
            return ResumeExtractor._extract_docx(file_path, file_bytes)

    @staticmethod
    def _extract_pdf(file_path: str = None, file_bytes: bytes = None) -> str:
        """Extract text from PDF."""
        if file_bytes:
            reader = PdfReader(io.BytesIO(file_bytes))
        else:
            reader = PdfReader(file_path)

        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        text = text.strip()
        if not text:
            raise EmptyResume("PDF appears to be empty or contains only images.")
        return text

    @staticmethod
    def _extract_docx(file_path: str = None, file_bytes: bytes = None) -> str:
        """Extract text from DOCX."""
        from docx import Document

        if file_bytes:
            doc = Document(io.BytesIO(file_bytes))
        else:
            doc = Document(file_path)

        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        if not text.strip():
            raise EmptyResume("DOCX appears to be empty.")
        return text


class ResumeAnalyzer:
    """Uses Gemini AI to parse resume text into a structured profile."""

    PARSE_PROMPT = """You are a resume parser. Analyze the following resume text and extract a structured profile.

Return ONLY valid JSON with this exact schema:
{{
    "name": "Full Name",
    "role": "Target role (e.g., Software Engineer)",
    "skills": ["skill1", "skill2", ...],
    "achievements": ["achievement1", "achievement2", ...],
    "email_context": "A concise 1-2 sentence summary of the candidate's background for use in outreach emails"
}}

Rules:
- Extract at most 5 achievements — pick the most impressive, quantified ones
- Skills should be specific technologies/tools, not soft skills
- email_context should be concise and highlight what makes this candidate stand out
- If information is missing, use empty string or empty array

Resume text:
---
{resume_text}
---"""

    def __init__(self, api_key: str, provider: str = "gemini"):
        """Initialize with a user's API key and provider name.

        Supported providers: gemini (default), groq, openai.
        """
        self.provider = provider
        self.api_key = api_key

    async def parse(self, resume_text: str) -> ResumeProfile:
        """Parse resume text into a structured profile using AI."""
        prompt = self.PARSE_PROMPT.format(resume_text=resume_text[:5000])

        if self.provider == "gemini":
            return await self._parse_gemini(prompt)
        elif self.provider in ("groq", "openai"):
            return await self._parse_openai_compat(prompt)
        else:
            return await self._parse_gemini(prompt)

    async def _parse_gemini(self, prompt: str) -> ResumeProfile:
        """Parse using Google Gemini via REST API (no SDK dependency)."""
        url = (
            "https://generativelanguage.googleapis.com/v1beta/models/"
            f"gemini-2.0-flash:generateContent?key={self.api_key}"
        )
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"responseMimeType": "application/json"},
        }
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(url, json=payload)
        if resp.status_code != 200:
            raise ValueError(f"Gemini API error {resp.status_code}: {resp.text[:300]}")
        data = resp.json()
        candidates = data.get("candidates", [])
        if not candidates:
            raise ValueError("Gemini returned no candidates")
        text = candidates[0]["content"]["parts"][0]["text"]
        return self._extract_profile(text)

    async def _parse_openai_compat(self, prompt: str) -> ResumeProfile:
        """Parse using Groq or OpenAI (both use the OpenAI-compatible REST API)."""
        base_url = (
            "https://api.groq.com/openai/v1" if self.provider == "groq"
            else "https://api.openai.com/v1"
        )
        model = (
            "llama-3.3-70b-versatile" if self.provider == "groq"
            else "gpt-4o-mini"
        )
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                f"{base_url}/chat/completions",
                headers={"Authorization": f"Bearer {self.api_key}"},
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.2,
                },
            )
            resp.raise_for_status()
            text = resp.json()["choices"][0]["message"]["content"]
        return self._extract_profile(text)

    def _extract_profile(self, raw: str) -> ResumeProfile:
        """Parse the JSON returned by any AI provider into a ResumeProfile."""
        raw = raw.strip()
        # Strip markdown code fences if present
        json_match = re.search(r"```json?\s*(.*?)\s*```", raw, re.DOTALL)
        if json_match:
            raw = json_match.group(1)

        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            return ResumeProfile(
                name="",
                role="",
                skills=[],
                achievements=[],
                email_context="",
            )

        return ResumeProfile(
            name=data.get("name", ""),
            role=data.get("role", ""),
            skills=data.get("skills", [])[:15],
            achievements=data.get("achievements", [])[:5],
            email_context=data.get("email_context", ""),
        )

    @staticmethod
    def _extract_name_fallback(text: str) -> str:
        """Simple fallback: first non-empty line is likely the name."""
        for line in text.strip().split("\n"):
            line = line.strip()
            if line and len(line.split()) <= 4 and not any(c.isdigit() for c in line):
                return line
        return ""


def save_resume_profile(user_id: str, profile: ResumeProfile, file_path: str) -> ResumeProfile:
    """Save parsed resume profile to database."""
    supabase = get_supabase_admin()

    data = {
        "user_id": user_id,
        "name": profile.name,
        "role": profile.role,
        "skills": profile.skills,
        "achievements": profile.achievements,
        "email_context": profile.email_context,
        "file_path": file_path,
    }

    # Upsert — one profile per user
    existing = supabase.table("resume_profiles").select("id").eq("user_id", user_id).execute()
    if existing.data:
        result = supabase.table("resume_profiles").update(data).eq("user_id", user_id).execute()
    else:
        result = supabase.table("resume_profiles").insert(data).execute()

    row = result.data[0]
    profile.id = row["id"]
    return profile


def get_resume_profile(user_id: str) -> Optional[ResumeProfile]:
    """Get saved resume profile for a user."""
    supabase = get_supabase_admin()
    result = supabase.table("resume_profiles").select("*").eq("user_id", user_id).execute()
    if not result.data:
        return None
    row = result.data[0]
    return ResumeProfile(
        id=row["id"],
        name=row["name"],
        role=row["role"],
        skills=row.get("skills", []),
        achievements=row.get("achievements", []),
        email_context=row.get("email_context", ""),
    )
